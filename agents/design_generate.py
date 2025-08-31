from __future__ import annotations
import base64, json, datetime, io, random
from typing import Dict, Any, Optional, List, Tuple
from shared.aws import bedrock_runtime
from shared.s3 import put_object
from shared.config import settings

def _vendor_from_model_id(model_id: str) -> str:
    mid = (model_id or "").lower()
    if mid.startswith("amazon.titan-image"):
        return "titan"
    if mid.startswith("stability.stable-diffusion"):
        return "sdxl"
    if "anthropic" in mid or "claude" in mid:
        return "anthropic"
    return "unknown"

def _payload_titan(prompt: str) -> Dict[str, Any]:
    return {
        "taskType": "TEXT_IMAGE",
        "textToImageParams": {"text": prompt},
        "imageGenerationConfig": {
            "numberOfImages": 1,
            "quality": "standard",
            "height": 1024,
            "width": 1024,
            "cfgScale": 8,
            "seed": 0,
        },
    }

def _payload_sdxl(prompt: str) -> Dict[str, Any]:
    return {
        "text_prompts": [{"text": prompt}],
        "cfg_scale": 8,
        "height": 1024,
        "width": 1024,
        "samples": 1,
        "steps": 30,
    }

def _placeholder_svg_bytes(title: str, subtitle: str) -> bytes:
    t = (title or "KaiKashi DreamForge").replace("&", "&amp;")
    s = (subtitle or "").replace("&", "&amp;")
    svg = f"""<?xml version="1.0"?>
<svg xmlns="http://www.w3.org/2000/svg" width="1024" height="1024">
  <defs>
    <linearGradient id="g" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0%" stop-color="#00e5ff"/>
      <stop offset="100%" stop-color="#7c4dff"/>
    </linearGradient>
  </defs>
  <rect width="100%" height="100%" fill="url(#g)"/>
  <g fill="white" font-family="Readex Pro, Inter, system-ui" text-anchor="middle">
    <text x="512" y="480" font-size="64" font-weight="700">{t}</text>
    <text x="512" y="540" font-size="28" opacity="0.85">{s}</text>
  </g>
</svg>"""
    return svg.encode("utf-8")

def _placeholder_obj_bytes(title: str) -> bytes:
    safe_title = (title or "kaikashi").replace("\n", " ")
    obj = f"""# KaiKashi placeholder
# title: {safe_title}
o plane
v -0.5 0.0 -0.5
v  0.5 0.0 -0.5
v  0.5 0.0  0.5
v -0.5 0.0  0.5
vt 0.0 0.0
vt 1.0 0.0
vt 1.0 1.0
vt 0.0 1.0
vn 0.0 1.0 0.0
usemtl default
s off
f 1/1/1 2/2/1 3/3/1
f 1/1/1 3/3/1 4/4/1
"""
    return obj.encode("utf-8")

def _make_book_outline(brief: Dict[str, Any]) -> List[str]:
    """
    Crea un bosquejo de capítulos a partir del intent/tags.
    Si ya tienes otra lógica/LLM para outline, puedes reemplazar aquí.
    """
    base = [
        "Introducción",
        "Contexto histórico",
        "Actores principales",
        "Eventos clave",
        "Impacto y consecuencias",
        "Conclusiones",
        "Bibliografía"
    ]
    # pimp simple por tags
    tags = [t.lower() for t in (brief.get("tags") or [])]
    if any("niñ" in t or "child" in t for t in tags):
        base.insert(2, "Glosario para jóvenes lectores")
    return base

def _build_book_docx_bytes(brief: Dict[str, Any], design_prompt: str) -> bytes:
    """
    Genera un DOCX de 'libro real' con:
    - Portada (título, subtítulo/estilo)
    - Índice (lista numerada de capítulos)
    - Introducción y capítulos con headings
    - Bibliografía (placeholder)
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

    title = brief.get("intent") or "Libro generado con KaiKashi"
    style = brief.get("style") or ""
    notes = brief.get("notes") or ""
    outline = _make_book_outline(brief)

    doc = Document()

    if "KaiKashi Body" not in [s.name for s in doc.styles]:
        body_style = doc.styles.add_style("KaiKashi Body", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = "Calibri"
        body_style.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(datetime.datetime.utcnow().strftime("%Y-%m-%d"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("Índice", level=1)
    for i, cap in enumerate(outline, 1):
        doc.add_paragraph(f"{i}. {cap}", style="List Number")
    doc.add_paragraph("Sugerencia: en Word use Referencias → Tabla de Contenido para insertar/actualizar el TOC.")

    doc.add_page_break()

    doc.add_heading("Introducción", level=1)
    intro = (
        f"Este libro aborda {title.lower()}. "
        f"Estilo: {style}. "
        f"Notas de producción: {notes}."
    )
    doc.add_paragraph(intro, style="KaiKashi Body")

    sample_paras = [
        "Este capítulo explora los antecedentes y las condiciones que dieron origen al tema. "
        "Se presentan líneas de tiempo, contextos geopolíticos y referencias comparativas.",

        "Se analizan las principales figuras y organizaciones involucradas, con atención a sus motivaciones, "
        "decisiones y consecuencias estratégicas.",

        "Se sintetizan los eventos clave de manera cronológica, incluyendo hitos, reacciones y repercusiones regionales.",

        "Se estudian los impactos sociales, económicos y culturales, así como lecciones aprendidas."
    ]
    for cap in outline:
        if cap.lower().startswith(("intro", "bibliograf")):
            continue
        doc.add_heading(cap, level=1)
        for _ in range(3):
            doc.add_paragraph(random.choice(sample_paras), style="KaiKashi Body")

    if not any("conclu" in c.lower() for c in outline):
        doc.add_heading("Conclusiones", level=1)
        doc.add_paragraph(
            "Resumen de hallazgos, líneas futuras de investigación y recomendaciones prácticas.",
            style="KaiKashi Body"
        )

    doc.add_heading("Bibliografía", level=1)
    bib_samples = [
        "Autor, A. (Año). Título del libro. Editorial.",
        "Autor, B. (Año). Artículo en Revista. Revista X, Vol(Y), pp–pp.",
        "Sitio/Institución (Año). Recurso en línea. URL."
    ]
    for b in bib_samples:
        doc.add_paragraph(f"- {b}", style="KaiKashi Body")

    doc.add_page_break()
    doc.add_heading("Anexo: Design Prompt", level=2)
    doc.add_paragraph(design_prompt, style="KaiKashi Body")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _build_book_txt_bytes(brief: Dict[str, Any], design_prompt: str) -> bytes:
    title = brief.get("intent") or "Libro"
    style = brief.get("style") or ""
    outline = _make_book_outline(brief)

    lines: List[str] = []
    lines += [title.upper(), "=" * len(title), "", f"Estilo: {style}", ""]
    lines += ["ÍNDICE", "------"]
    for i, cap in enumerate(outline, 1):
        lines.append(f"{i}. {cap}")
    lines += ["", "INTRODUCCIÓN", "------------",
              f"Este libro aborda {title.lower()}.",
              ""]
    for cap in outline:
        if cap.lower().startswith(("intro", "bibliograf")):
            continue
        lines += [cap.upper(), "-" * len(cap),
                  "Contenido de capítulo (placeholder).", ""]
    lines += ["CONCLUSIONES", "------------", "Resumen de hallazgos y recomendaciones.", ""]
    lines += ["BIBLIOGRAFÍA", "------------",
              "- Autor, A. (Año). Título del libro. Editorial.",
              "- Autor, B. (Año). Artículo en Revista. Revista X, Vol(Y), pp–pp.",
              "- Sitio/Institución (Año). Recurso en línea. URL.", "",
              "ANEXO: DESIGN PROMPT", "-------------------", design_prompt]
    return ("\n".join(lines)).encode("utf-8")

def _decide_kinds(brief: Dict[str, Any]) -> List[str]:
    """
    Decide los tipos a generar. ¡Sin PDF!
    """
    pt = (brief.get("product_type") or "").lower()
    intent = (brief.get("intent") or "").lower()

    img_only = {"poster", "tshirt", "sticker", "children_book_cover",
                "book_cover", "cover", "ebook_cover", "mockup", "image"}
    doc_only = {"book", "ebook", "educational_content", "story",
                "children_book", "novel", "document", "libro"}
    vid_only = {"video", "clip", "animation"}
    model_only = {"3d_model", "3d_printable", "model3d", "3d"}

    if pt in img_only or "image" in intent:
        return ["image"]
    if pt in doc_only or ("book" in intent and pt == "") or ("libro" in intent and pt == ""):
        # Sin PDF: libro real = DOCX + TXT
        return ["docx", "txt"]
    if pt in vid_only or "video" in pt or "video" in intent or "gif" in intent:
        return ["video"]
    if pt in model_only or "3d" in pt or "3d" in intent:
        return ["3d"]
    return ["image"]

def generate_assets(design_prompt: str, brief: Dict[str, Any], user_id: str) -> Dict[str, Any]:
    outputs: Dict[str, Any] = {}
    model_id = getattr(settings, "bedrock_image_model_id", "")
    vendor = _vendor_from_model_id(model_id)
    base = f"assets/{user_id}/generated/{brief.get('product_type','generic')}_{brief.get('intent','idea')}"
    errors: Dict[str, str] = {}
    media_keys: List[str] = []
    kinds = _decide_kinds(brief)

    # Image
    if "image" in kinds:
        image_key: Optional[str] = None
        try:
            if vendor in ("titan", "sdxl"):
                rt = bedrock_runtime()
                body = _payload_titan(design_prompt) if vendor == "titan" else _payload_sdxl(design_prompt)
                res = rt.invoke_model(modelId=model_id, body=json.dumps(body))
                payload = json.loads(res["body"].read())
                if vendor == "titan":
                    img_b64 = (payload.get("images") or [None])[0] or payload.get("image_base64")
                else:
                    artifacts = payload.get("artifacts", [])
                    img_b64 = artifacts[0].get("base64") if artifacts else None
                if img_b64:
                    raw = base64.b64decode(img_b64)
                    image_key = f"{base}.png"
                    put_object(settings.s3_bucket_assets, image_key, raw, "image/png")
                else:
                    errors["image"] = "Modelo de imagen no devolvió salida base64."
            elif vendor == "anthropic":
                errors["image"] = "El modelo configurado es Anthropic/Claude (no genera imágenes)."
            else:
                errors["image"] = f"Modelo '{model_id}' no reconocido como generador de imagen."
        except Exception as e:
            errors["image"] = f"{type(e).__name__}: {e}"

        if not image_key:
            title = brief.get("intent") or "Diseño generado"
            subtitle = (brief.get("style") or "")[:80]
            svg_bytes = _placeholder_svg_bytes(title, subtitle)
            ts = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
            image_key = f"{base}_placeholder_{ts}.svg"
            put_object(settings.s3_bucket_assets, image_key, svg_bytes, "image/svg+xml")

        outputs["image_key"] = image_key
        media_keys.append(image_key)

    # (DOCX + TXT)
    if "docx" in kinds:
        try:
            docx_bytes = _build_book_docx_bytes(brief, design_prompt)
            docx_key = f"{base}.docx"
            put_object(settings.s3_bucket_assets, docx_key, docx_bytes,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            outputs["docx_key"] = docx_key
            media_keys.append(docx_key)
        except Exception as e:
            errors["doc"] = f"{type(e).__name__}: {e}"

    if "txt" in kinds:
        try:
            txt_bytes = _build_book_txt_bytes(brief, design_prompt)
            txt_key = f"{base}.txt"
            put_object(settings.s3_bucket_assets, txt_key, txt_bytes, "text/plain; charset=utf-8")
            outputs["text_key"] = txt_key
            media_keys.append(txt_key)
        except Exception as e:
            errors["text"] = f"{type(e).__name__}: {e}"

    # Video (GIF placeholder)
    if "video" in kinds:
        try:
            from PIL import Image, ImageDraw
            imgs: List[Any] = []
            for i in range(12):
                img = Image.new("RGB", (1024, 1024))
                d = ImageDraw.Draw(img)
                d.rectangle((0, 0, 1024, 1024), fill=(10, 10, 20))
                d.ellipse((112+i*2, 112, 912, 912), fill=(10, 150, 230))
                d.ellipse((212, 212, 812-i*2, 812), fill=(120, 80, 255))
                imgs.append(img)
            buf = io.BytesIO()
            imgs[0].save(buf, format="GIF", save_all=True, append_images=imgs[1:], duration=80, loop=0)
            gif_key = f"{base}.gif"
            put_object(settings.s3_bucket_assets, gif_key, buf.getvalue(), "image/gif")
            outputs["video_key"] = gif_key
            media_keys.append(gif_key)
        except Exception as e:
            errors["video"] = f"{type(e).__name__}: {e}"

    # 3D placeholder
    if "3d" in kinds:
        try:
            obj_bytes = _placeholder_obj_bytes(brief.get("intent", ""))
            obj_key = f"{base}.obj"
            put_object(settings.s3_bucket_assets, obj_key, obj_bytes, "text/plain")
            outputs["model3d_key"] = obj_key
            media_keys.append(obj_key)
        except Exception as e:
            errors["3d"] = f"{type(e).__name__}: {e}"

    outputs["kinds"] = kinds
    outputs["media_keys"] = media_keys
    outputs["package"] = {
        "design_prompt": design_prompt,
        "brief": brief,
        "suggested_title": brief.get("intent", "Producto creativo"),
        "suggested_description": f"Generado automáticamente a partir de tu idea: {brief.get('notes','')}",
    }
    if errors:
        outputs["errors"] = errors
    return outputs
